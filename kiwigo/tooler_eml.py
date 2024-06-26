'''
本程序用于；
   - 解析.eml，并提取授权单（一种固定格式的.docx文件）
   - 对提取出的授权单按照一定顺序编号（结合邮件到达时间，优先以同一封邮件多个授权单连号的形式分配授权单编号）

v1.01 更新内容：
   - 修复了已归档授权单为空时报错
   - 增加了自动修复命名不规范的归档授权单功能（尽可能修复，无法修复则需要手动修改）
v1.02 更新内容：
   - 修复了邮件读取目录的笔误
   - 修复了已识别授权单为纯连号或没有最大号内不够插空分配新单号的bug
   - 修复了docx内重写标题（授权单号）后未保存的bug
   - 增加了识别两页以上授权单转pdf转png的功能（*但是多页授权单表格是否能正常读取仍需观察）
v1.03 更新内容：
   - 修复了for循环中try报错break直接跳出循环的错误，改正为用continue跳出错误继续完成循环
   - 增加了.bat自动复制“【授权】”等字样的头信息
   - 区分科运中心、软件中心的机房授权单
   - 加入输出文件夹最大buffer，超过最大数后运行程序会自动清理
   - 增加【授权主题】替换一些无效字符（如Fw:）
v23.45.0 更新内容：
   - 增加了[和林格尔]新园区机房
   - 更新了[XX机房授权]
   - 优化了部分输出排版
   - 简单实现删除解析正常的.eml
v23.46.0 更新内容：
   - 增加了授权文件按自然排序分配id（即《授权单10.docx》在《授权单6.docx》之后分配编号）
   - 增加了剪切到工作目录并解压筛选功能（【提示】使用本程序解压.rar文件前，需先在环境变量中添加winrar.exe所在根目录，并重启pycharm）
   - 修复了一些bug（如复制发布时同邮件授权单号不按顺序拍）
v24.0000
   - 增加.eml归档功能（以授权单命名）
v24.21.0
   - kgo.emls_to_doxs()方法新增参数sender_mode='forward'/'origin'，'forward'模式下可识别申请人过指定人转发来的邮件，并视申请人为发件人
   - 放宽机房名过滤要求

todo:
   - 增加.xlsx文件全面记录信息
   - 简化变量命名
   - 进一步解决无归档授权单时报错的算法结构（虽然目前已能初步debug，可以使用set）
   - 考虑是否增加可以自动初始化工作目录的方法（比如新建导入的文件路径）
   - 系统解决.bat转义字符（如&等）
   - fix_sqd()无法识别.~前缀的.docx文件
   - 增加log
   - 归档文件命名修复（如A7-405无法识别）
   - 增减单号
   - 自动除空格
   - 尽可能正确识别各种形式日期，规范化

questions:
   - fitz.open()报黄

cautions:
   - 最大归档授权单号不得超过10000

@author: tyhua
'''

import eml_parser
import os
import base64
import datetime
import time
import re
import shutil
import pandas as pd
import copy

import natsort
import pyperclip
import docx
from docx.enum.style import WD_STYLE_TYPE
from docx import Document
from colorama import Fore, Back, Style
import comtypes.client
import fitz

import pprint

from .toolman import Toolman

tm = Toolman()
_dct_ini_ = tm._rd_ini()

_path_ = _dct_ini_['_dir_run_py_']

# 用于存放已归档的授权文件
_dir_afbase_ = _dct_ini_['_dir_afbase_']
_dir_afb_kyzx_t_ = _dct_ini_['_dir_afb_kyzx_t_']
_dir_afb_kyzx_u_ = _dct_ini_['_dir_afb_kyzx_u_']
_dir_afb_rjzx_ = _dct_ini_['_dir_afb_rjzx_']
_dir_afb_hlge_ = _dct_ini_['_dir_afb_hlge_']
_dir_eml_ = _dct_ini_['_dir_eml_']
_out_path_hd_ = _dct_ini_['_hd_dir_out_']
_out_path_ = _dct_ini_['_dir_out_']
# 用于存放识别成功的授权单文件（.docx）
_dir_docx_ = os.path.join(_out_path_, '#_0_docx_#')
_bsname_xlsx_ = 'a.xlsx'
_path_xlsx_ = os.path.join(_path_, 'a.xlsx')

_lst_kyzx_t_ = ['A3-302', 'A3-305', 'A3-402', 'A3-405', 'A7-405', 'A7-406']
_lst_kyzx_u_ = ['DC1-201', 'DC1-202', 'DC1-301', 'DC2-203', 'DC2-401', ]
_lst_rjzx_ = ['A3-303']
_lst_hlge_ = ['B1-1A','B1-1B','B1-2A','B1-2B','B1-2C','B1-2D','B1-3A','B1-3B','B1-3C','B1-3D','B3-1A','B3-1B','B3-2A','B3-2B','B3-2C','B3-2D','B3-3A','B3-3B','B3-3C','B3-3D']

_dct_pak_rom = {
    '电信': _lst_kyzx_t_,
    '联通': _lst_kyzx_u_,
    '软件中心': _lst_rjzx_,
    '和林格尔': _lst_hlge_
}

# 用于保存邮件信息
_mail_txt_ = '#_mail_#.txt'
_bat_cmd_ = 'echo ☆ |clip'
_bat_fn_ = '##★_双击复制【授权主题】##.bat'

# 生成【授权主题】时，屏蔽的词
_lst_ban_wd_ = ['Fw:', 'Fw', '以此为准', '[', ']', '【', '】', '转发', '{', '}']

# 生成表信息
_tab_hdlst_ = ['申请来源', '机房授权表编号', '机房', '人员所属', '事由', '日期', '整机', '配件', '变更', '授权人数', '实际人数']
_tab_pth_ = os.path.join(_out_path_, '#_1_xlsx.xlsx')

_tab_hdlst_ = ['序号', '申请来源', '机房授权编号', '机房', '人员所属', '变更号', '进出时间', '授权人数', '实际人数']

_lst_cps_ = ['.zip', '.rar',]   #  '.7z'的包离线太难装了

# 变更中需要重复去除的符号
_lst_sig_ = [' ', ';', '；', '&', '和', ',', '，', '\\', '/', '+', '|', '、', '\xa0', '\t', '\n']
_lst_docx_sig_ = ['起始日期', '截止日期', '事由', '人员', '设备', '变更单号', '设备名称', '设备数量']

# log
_lst_hd_log_ea = ['is_done', 'date', 'sender', 'dis_a_zone', 'dis_a_id', 'room', 'sbj', 'af/afs', 'af_org', 'af_edt', 'pth_eml', 'dir_eml_i']


class Sq:
    def __init__(self):
        self.tm = Toolman()
        self.dict_syn = {}

        # 用于记录错误信息
        self.dict_log = {}
        self.dict_log['er_eml'] = []
        self.dict_log['dir'] = []   # 记录输出目录相关信息
        self.dict_log['ea'] = {}   # 记录输出邮件&授权单相关信息

        self.__dir_init(_dir=_dir_eml_)

    def __dir_init(self, _dir):
        if not os.path.exists(_dir):
            os.makedirs(_dir)

    # 对以归档授权单进行处理
    @classmethod
    def _fix_sqd(cls, _dir_afb):
        # print('-'*10, '[start]', _dir_afb, '-'*10)
        lst_bad = []
        lst_bad2 = []
        lst_docx = []
        lst_id = []
        lst_idfed = []

        for _dir in os.listdir(_dir_afb):
            if _dir.endswith('.docx'):
                lst_docx.append(_dir)
                afn_i, _ = os.path.splitext(_dir)

                # 一次检验
                if not re.search(
                        '^20\d\d[0-1]\d-\d{2,3}-\w{3,5}(\()+\d{4}-\d{1,2}-\d{1,2} [0-9]{6}[\u4e00-\u9fa5]{2,3}(\))+$',
                        afn_i):
                    lst_bad.append(afn_i)
                else:
                    lst_id.append(int(re.findall('-\d*-', afn_i)[0][1:-1]))
                    lst_idfed.append(afn_i)

        for afn_i in lst_bad:
            a0 = afn_i

            # 若满足最宽泛的条件才有修改的可能，否则无法修改，只能看一眼手动去改
            if re.search(
                    '^[\s]*[0-9]{6}[\s]*-[\s]*[0-9]{2,3}[\s]*-[\s]*.{3,5}[\s]*(\(|[\uff08])+[\s]*\d{4}[\s]*-[\s]*\d{1,2}[\s]*-[\s]*\d{1,2}[\s]*[0-9]{6}[\s]*[\u4e00-\u9fa5]{2,3}[\s]*(\)|[\uff09])[\s]*$',
                    afn_i):
                # [修复情况1]去除多余空格
                afn_i = afn_i.replace(' ', '')
                pointer_blk = [i.span() for i in re.finditer('[0-2]\d[0-6]\d[0-6]\d[\u4e00-\u9fa5]', afn_i)][0][0]
                l_ = list(afn_i)
                l_.insert(pointer_blk, ' ')
                afn_i = ''.join(l_)

                # [修复情况2]括号切换中/英文
                # print(afn_i)
                afn_i = afn_i.replace('（', '(')
                afn_i = afn_i.replace('）', ')')

            # 复检
            if re.search(
                    '^20\d\d[0-1]\d-\d{2,3}-\w{3,5}(\()+\d{4}-\d{1,2}-\d{1,2} [0-9]{6}[\u4e00-\u9fa5]{2,3}(\))+$',
                    afn_i):
                os.rename(os.path.join(_dir_afb, a0 + '.docx'), os.path.join(_dir_afb, afn_i + '.docx'))
                print(os.path.split(_dir_afb)[-1], '修复完成：', a0, '\t->\t', afn_i)
            else:
                lst_bad2.append(afn_i)

        print(os.path.split(_dir_afb)[-1], '授权单命名仍未修复：', len(lst_bad2), ' ', lst_bad2) if lst_bad2 else print(
            os.path.split(_dir_afb)[-1], '目录下授权单命名检验正常')

        # 输出
        print('识别到授权单号（个） - 识别到授权文件（个） / 共有.docx文件（个）: \t', len(set(lst_id)), '-', len(lst_idfed), '/', len(lst_docx))
        lst_id.sort()

        # 1. 查空缺
        lst_tmp1 = list(set(lst_id))
        nums_id = len(lst_tmp1)
        while True:
            id_ = cls.__dis_aid(lst_dised=lst_tmp1, lst_x=[1])
            if id_[0][0] > nums_id:
                break
            else:
                lst_tmp1.append(id_[0][0])
                print('[查到此编号空缺]', id_[0][0])
        # 2. 查多余
        lst_tmp2 = lst_id
        for i in set(lst_tmp2):
            lst_tmp2.remove(i)
            if i in lst_tmp2:
                print('[查到此编号多余]', i)

        # print('-' * 10, '[end]', _dir_afb, '-' * 10, '\n')

    def fix_sqd(self, ):
        print('\n' + '-' * 10, '[start]\t已归档授权单检查', '-' * 10)
        Sq._fix_sqd(_dir_afb=_dir_afb_kyzx_t_)
        Sq._fix_sqd(_dir_afb=_dir_afb_kyzx_u_)
        Sq._fix_sqd(_dir_afb=_dir_afb_rjzx_)
        Sq._fix_sqd(_dir_afb=_dir_afb_hlge_)
        print('-' * 10, '[end]\t已归档授权单检查', '-' * 10, '\n')

    def _chg_aid(self, _docx, dict_syn_ki):
        try:
            d0 = Document(_docx)
            t_ = d0.tables  # 获取文件中的表格集
            t0 = t_[0]
        except:
            return print(_docx, '改aid失败')
        self.tm.fun_chg_run_text(t1=t0.cell(0, 1),
                                 _txt='人员设备进出机房授权表（编号' + re.findall('^[0-9]{6}-[0-9]{2,3}', dict_syn_ki['name_nw'])[
                                     0] + '）')
        d0.save(_docx)

    def _idf_sqd(self, eml_idf, dir_eml_atm, dir_eml):
        # for f_i in os.listdir(dir_eml_atm):  # 仅遍历当前文件，不穿透深层文件夹
        # if f_i.endswith('.docx'):
        #     print(f_i)
        # a = os.path.join(dir_eml_atm, '202307-178（A3302 2023-07-26 256161 滑天扬）.docx')
        if not os.path.exists(_dir_docx_):
            os.makedirs(_dir_docx_)

        # 1. get
        docx_list = []
        dict_docx = {}

        for _dir in os.listdir(dir_eml_atm):
            if _dir.endswith('.docx'):
                docx_list.append(_dir)

        if not docx_list:
            dict_docx[-1] = {}  # 代表此邮件无授权单，否则数值代表其中有几个授权单

        # 2. output
        list_bad = []
        list_good = []
        list_af_docx = []

        # 判断授权文件，获取授权文件list
        for af_i in docx_list:
            try:
                d0 = Document(os.path.join(dir_eml_atm, af_i))
                t_ = d0.tables  # 获取文件中的表格集
                t0 = t_[0]
                t_in = self.tm.get_nested_tables_solu1(t0)
                list_af_docx.append(af_i)
            except:
                # print('[不是授权单，跳过]', af_i)
                continue

        list_af_docx = natsort.natsorted(list_af_docx)  # 排序，但不足的是无法和windows下文件排序一致（符号无法排在数字前）

        # 读取授权文件
        for _id, af_i in enumerate(list_af_docx):
            d0 = Document(os.path.join(dir_eml_atm, af_i))
            t_ = d0.tables  # 获取文件中的表格集
            t0 = t_[0]
            t_in = self.tm.get_nested_tables_solu1(t0)

            afn_i, _ = os.path.splitext(af_i)
            dict_docx[_id] = {}
            # if fun_get_run_text(t1=t0.cell(8, 1).paragraphs[0]):

            dict_docx[_id]['af_name'] = afn_i
            dict_docx[_id]['进出时间'] = self.tm.fun_get_run_text(t1=t0.cell(8, 1)) + \
                                     ' ~ ' + self.tm.fun_get_run_text(t1=t0.cell(8, 6))
            dict_docx[_id]['wkdate'] = self.tm.fun_get_run_text(t1=t0.cell(8, 1))

            _room = self.tm.fun_get_run_text(t1=t_in.cell(1, 2))
            _room = tm.rm_sig(_in=_room, _rm=' ', del_blk=True)  # delete the blank str
            dict_docx[_id]['room'] = self._identify_room(_r_str=_room)
            # todo: 方法加入mode::easy\hard
            # todo: 后续增加功能，将正确机房写入授权单中

            try:
                dict_docx[_id]['af_zone'] = \
                    [ki for ki in _dct_pak_rom.keys() if dict_docx[_id]['room'] in _dct_pak_rom[ki]][0]
            except:
                print(Fore.WHITE + Back.BLACK + Style.DIM + '[error] ' + os.path.join(dir_eml_atm,
                                                                                      dict_docx[_id]['af_name'])
                      + ' 或因未识别成功机房\t未找到读取到的机房：' + _room + Style.RESET_ALL)
                self.dict_log['er_eml'].append(dir_eml)
                dict_docx[-1] = {}

            if self.tm.fun_get_run_text(t1=t_in.cell(0, 1)):
                # print('入场')
                dict_docx[_id]['入出场'] = '入场'
            elif self.tm.fun_get_run_text(t1=t_in.cell(0, 3)):
                # print('出场')
                dict_docx[_id]['入出场'] = '出场'
            else:
                list_bad.append(af_i)
                break

            dict_docx[_id]['厂商'] = self.tm.fun_get_run_text(t1=t0.cell(2, 2)) + ';\n' + self.tm.fun_get_run_text(
                t1=t0.cell(1, 2))
            dict_docx[_id]['厂商人数'] = sum(
                [int(i) for i in re.findall('[0-9]', self.tm.fun_get_run_text(t1=t0.cell(2, 2)))])

            str1 = self.tm.fun_get_run_text(t1=t0.cell(3, 2)) + '×' + self.tm.fun_get_run_text(t1=t0.cell(4, 2)) + ';' \
                   + self.tm.fun_get_run_text(t1=t0.cell(3, 4)) + '×' + self.tm.fun_get_run_text(t1=t0.cell(4, 4)) + ';' \
                   + self.tm.fun_get_run_text(t1=t0.cell(3, 7)) + '×' + self.tm.fun_get_run_text(t1=t0.cell(4, 7))
            str1 = str1.replace(';×', '')
            if str1.endswith('×'): str1 = str1.replace('×', '')
            if str1.endswith(';'): str1 = str1.replace(';', '')
            dict_docx[_id]['设备'] = str1

            # CHG
            dict_chg = {}
            for i in range(10):
                dict_chg[_id] = []
                dict_chg[_id].append(self.tm.fun_get_run_text(t1=t0.cell(6, i)))
                dict_chg['num_' + str(_id)] = [i.count('CHG') for i in dict_chg[_id]]

            dict_docx[_id]['变更标记（与生产管理核对）'] = dict_chg['num_' + str(_id)][0] \
                if dict_chg['num_' + str(_id)][0] > 0 else ''
            dict_docx[_id]['变更号'] = dict_chg[_id]

            # 申请具体事项
            # df0.loc[_id, '申请具体事项'] = '="text1"&CHAR(10)&"text2"'
            dict_docx[_id]['申请具体事项'] = str(dict_docx[_id]['厂商']) \
                                       + '\n' + str(dict_docx[_id]['变更号']) \
                                       + '\n' + str(self.tm.fun_get_run_text(t1=t0.cell(7, 1)))

            list_good.append(af_i)

        print('邮件中\t授权单数 / docx文件数: ', len(list_good), ' /', len(docx_list))
        if list_bad:
            print('邮件中\t未识别为授权单的docx文件: ', len(list_bad), ', 是：', list_bad)

        # 5. 转移授权单docx
        print(_dir_docx_, dict_docx)
        for af_i in list_good:
            shutil.copyfile(src=os.path.join(dir_eml_atm, af_i), dst=os.path.join(_dir_docx_, eml_idf + af_i))

        return dict_docx

    # 识别机房
    def _identify_room(self, _r_str):
        # 若此处报错无['sbj_aid']或因未识别成功机房引起
        # ↓此法仅严格识别
        # dict_docx[_id]['af_zone'] = \
        # [ki for ki in _dct_pak_rom.keys() if dict_docx[_id]['room'] in _dct_pak_rom[ki]][0]
        '''放宽机房识别条件'''
        for k in _dct_pak_rom.keys():
            for _r in _dct_pak_rom[k]:
                if _r_str == _r:
                    return _r
                elif re.sub(r'[^a-zA-Z0-9\u4e00-\u9fa5\s]+', '', _r_str).upper() \
                        == re.sub(r'[^a-zA-Z0-9\u4e00-\u9fa5\s]+', '', _r).upper().upper():
                    print('[提示] 授权单机房可能填写不规范（已由' + _r_str + '模糊识别为' + _r + '，可忽略）')
                    return _r
        return -1

    def _sqd_png(self, _from_dx, _to_dir):
        _fm_dir, _fm_dx = os.path.split(_from_dx)
        _dx, _ = os.path.splitext(_fm_dx)
        dir_pdf = os.path.join(_fm_dir, _dx + '.pdf')
        pth_png = '_' + _dx

        # 1. .docx -> .pdf
        word = comtypes.client.CreateObject('Word.Application')
        doc = word.Documents.Open(_from_dx)
        doc.SaveAs(dir_pdf, FileFormat=17)
        doc.Close()
        word.Quit()

        # 2. .pdf -> .png
        d = fitz.open(dir_pdf)  # open document
        for _pg in range(0, d.page_count - 1):
            page_ = d[_pg]
            pix = page_.get_pixmap()  # render page to an image
            pix.save(os.path.join(_to_dir, '#+' + pth_png + str(_pg) + '+#.png'))  # store image as a PNG
        d.close()

        # 3. del
        os.remove(dir_pdf)
        print('[删除文件] ', dir_pdf)

    def _syn_dict(self, dict_docx, dict_eml):
        dict_docx = dict_docx
        dict_eml = dict_eml
        dict_syn = self.dict_syn

        for _id in dict_docx.keys():
            d_key = str(str(dict_eml['date']) + dict_eml['sender'] + str(_id)).replace(' ', '')

            dict_syn[d_key] = {}
            dict_syn[d_key]['e_id'] = str(_id + 1) + '/' + str(len(dict_docx.keys()))
            dict_syn[d_key]['date'] = dict_eml['date']  # todo: use datetime格式
            dict_syn[d_key]['sender'] = dict_eml['sender']
            dict_syn[d_key]['pth_eml'] = dict_eml['pth_eml']
            dict_syn[d_key]['dir_attach'] = dict_eml['dir_attach']
            dict_syn[d_key]['sbj'] = dict_eml['sbj']
            dict_syn[d_key]['eml_idf'] = dict_eml['eml_idf']

            dict_syn[d_key]['af_name'] = dict_docx[_id]['af_name']
            dict_syn[d_key]['wkdate'] = dict_docx[_id]['wkdate']
            dict_syn[d_key]['af_zone'] = dict_docx[_id]['af_zone']
            dict_syn[d_key]['room'] = dict_docx[_id]['room']

            # for log
            key2 = str(dict_eml['date']) + '_' + dict_eml['sender']
            key3 = int(str(_id).replace(' ', ''))
            self.dict_log['ea'][key2][key3] = {}   # 每个授权单的dict

            self.dict_log['ea'][key2][key3]['e_id'] = str(_id + 1) + '/' + str(len(dict_docx.keys()))
            self.dict_log['ea'][key2][key3]['date'] = dict_eml['date']  # todo: use datetime格式
            self.dict_log['ea'][key2][key3]['sender'] = dict_eml['sender']
            self.dict_log['ea'][key2][key3]['pth_eml'] = dict_eml['pth_eml']
            self.dict_log['ea'][key2][key3]['dir_attach'] = dict_eml['dir_attach']
            self.dict_log['ea'][key2][key3]['sbj'] = dict_eml['sbj']
            self.dict_log['ea'][key2][key3]['eml_idf'] = dict_eml['eml_idf']

            self.dict_log['ea'][key2][key3]['af_name'] = dict_docx[_id]['af_name']
            self.dict_log['ea'][key2][key3]['wkdate'] = dict_docx[_id]['wkdate']
            self.dict_log['ea'][key2][key3]['af_zone'] = dict_docx[_id]['af_zone']
            self.dict_log['ea'][key2][key3]['room'] = dict_docx[_id]['room']

        self.dict_syn = dict_syn

    # 输入一个list（def内自动排序），输出不连续空出的单号
    @classmethod
    def __out_blk(cls, lst):
        lst = sorted(lst)
        lst_blk = []
        for _p, i in enumerate(lst):
            if _p == len(lst) - 1:
                break
            blk = lst[_p + 1] - lst[_p] - 1
            lst_blk.append(blk)
        return lst_blk

    @classmethod
    def __dis_aid(cls, lst_dised, lst_x, ):
        # lst_dised = [2, 3, 5, 12]   # 已经分配过的单号
        # lst_x = [3, 1, 2]   # 有几个单子待分配单号
        if not 0 in lst_dised:
            lst_dised.append(0)
        # for debug
        lst_dised.append(10000)
        # todo: use an elegant way to slove the problem, not just for debug
        lst_dised.sort()
        # print(lst_dised, type(lst_dised[0]), type(lst_dised[1]))
        lst_aid = []
        for k in lst_x:
            lst_blk = cls.__out_blk(lst=lst_dised)
            for _b, b in enumerate(lst_blk):
                if k <= b:
                    aid_ = list(range(lst_dised[_b] + 1, lst_dised[_b] + k + 1))
                    lst_aid.append(aid_)
                    lst_dised[_b + 1: _b + 1] = iter(aid_)
                    break
        return lst_aid

    def _dist_afid(self, dir_afb, dict_syn, lst_k):
        # 1. 收集目录下授权文件，返回一个list
        # DC1-202 DC1-301 DC1-201 DC2-203 DC2-401

        lst_auth_id = []
        file_cannot_identify_list = []
        for _dir in os.listdir(dir_afb):
            if _dir.endswith('.docx'):
                if re.search(
                        '^[\s]*[0-9]{6}-[0-9]{2,3}-.{3,5}(\(|[\uff08])+\d{4}-\d{1,2}-\d{1,2} [0-9]{6}[\u4e00-\u9fa5]{2,3}(\)|[\uff09])+(.docx)$',
                        _dir):
                    auth_mth_i = int(_dir.split('-')[0].replace(' ', ''))
                    auth_id_i = int(_dir.split('-')[1].replace(' ', ''))  # int
                    lst_auth_id.append(auth_id_i)
                else:
                    file_cannot_identify_list.append(_dir)

        # 2. 分派
        lst_bk = []
        lst_k = natsort.natsorted(lst_k)
        lst01 = []
        for ki in lst_k:
            lst01.append(ki[:16])

        lst02 = list(set(lst01))
        lst02.sort()

        for s in lst02:
            lst_bk.append(lst01.count(s))

        # 若未识别到本区域授权单，则return
        if lst_bk:
            print('-' * 10, '↓', '授权单处理', os.path.split(dir_afb)[-1], '↓', '-' * 10)
            print(os.path.split(dir_afb)[-1], '待分配单号的单数：', lst_bk)
        else:
            return 0

        print(os.path.split(dir_afb)[-1], '已识别到归档授权编号：', sorted(lst_auth_id))
        if file_cannot_identify_list:
            print(os.path.split(dir_afb)[-1], '以下文件无法识别为已授权文件：', file_cannot_identify_list)

        # 2. 分配af单号
        # todo: 暂时无法识别邮件中自带顺序的授权单 [natsort部分解决了该问题]
        lst_aid = self.__dis_aid(lst_dised=lst_auth_id, lst_x=lst_bk)  # lst_dised: int
        print(os.path.split(dir_afb)[-1], '已分配授权单号：', lst_aid)

        # 3. 复写
        lst_aid = [i for p in lst_aid for i in p]
        for _i, ki in enumerate(lst_k):
            dict_syn[ki]['a0_id'] = str(int(dict_syn[ki]['e_id'].split('/')[0]) - 1)   # 因为是按 1, 2, 3, ..., 100 这样排的
            dict_syn[ki]['name_nw'] = str(dict_syn[ki]['name_nw']).replace('X', str(lst_aid[_i]).rjust(2, '0'))
            dict_syn[ki]['sbj_aid'] = re.findall('\d{6}-\d{2,3}', dict_syn[ki]['name_nw'])[0][4:]

            self._chg_aid(_docx=os.path.join(_dir_docx_,dict_syn[ki]['eml_idf'] + dict_syn[ki]['af_name'] + '.docx'), dict_syn_ki=dict_syn[ki])

            # 源文件改名
            os.rename(os.path.join(_dir_docx_, dict_syn[ki]['eml_idf'] + dict_syn[ki]['af_name'] + '.docx'),
                      os.path.join(_dir_docx_, dict_syn[ki]['name_nw'] + '.docx'))
            print('已重命名：\t', ki, '\t', dict_syn[ki]['af_name'], '\tto\t', dict_syn[ki]['name_nw'])

            k1 = dict_syn[ki]['eml_idf']
            k2 = int(dict_syn[ki]['a0_id'])
            self.dict_log['ea'][k1][k2]['a0_id'] = k2
            self.dict_log['ea'][k1][k2]['name_nw'] = dict_syn[ki]['name_nw']
            self.dict_log['ea'][k1][k2]['sbj_aid'] = dict_syn[ki]['sbj_aid']

            # 截图授权单
            self._sqd_png(_to_dir=os.path.join(dict_syn[ki]['dir_attach']),
                          _from_dx=os.path.join(_dir_docx_, dict_syn[ki]['name_nw'] + '.docx'))

        return 0

        # 4. 生成信息
        # todo: 生成一个表记录信息，便于出错排障

    def _name_sqd(self, dict_syn):
        for d_key in dict_syn.keys():
            dict_syn[d_key]['wkdate'] \
                = tm.rm_sig(_in=dict_syn[d_key]['wkdate'], _rm=' ', del_blk=True)   # delete the blank str
            set_bad = set()
            # 1. 识别施工日期
            # todo: 有没有更简洁的替代方法
            lst_ymd = ['%Y-%m-%d', '%Y%m%d', '%Y.%m.%d', '%Y年%m月%d日']
            for i in lst_ymd:
                try:
                    wkdate = datetime.datetime.strptime(dict_syn[d_key]['wkdate'], i)
                    dict_syn[d_key]['wkdate_r'] = datetime.datetime.strftime(wkdate, '%Y%m')
                    break
                except:
                    set_bad.add(dict_syn[d_key]['af_name'])

            date_d = datetime.datetime.strptime(dict_syn[d_key]['date'], '%Y%m%d%H%M%S')
            date_ymr = datetime.datetime.strftime(date_d, '%Y-%m-%d')
            date_hms = datetime.datetime.strftime(date_d, '%H%M%S')

            d_t = dict_syn[d_key]
            d_t['name_nw'] = d_t['wkdate_r'] + '-X-' + d_t['room'].replace('-', '') + '(' + date_ymr + ' ' + date_hms + \
                             d_t['sender'] + ')'

        return dict_syn

    # todo: 检查dict_syn
    def _chk_dct(self, dict_syn):
        print(dict_syn)

        lst_bad = []
        lst_len_k = []

        for ki in dict_syn.keys():
            lst_len_k.append(len(dict_syn[ki].keys()))

        m_ = max(lst_len_k)

    def _txt(self, dict_syn):
        # 1. 按邮件为单位分类
        lst_1 = []
        for ki in dict_syn.keys():
            dict_syn[ki]['af_k'] = ki
            lst_1.append(dict_syn[ki])

        dct_1 = {}
        for i in lst_1:
            dir_bn = os.path.basename(i['dir_attach'])
            if not dir_bn in dct_1.keys():
                dct_1[dir_bn] = {}
            dct_1[dir_bn][i['af_k']] = i

        # todo: 简化以下处理过程的复杂度

        # 按邮件为单位，处理
        dct_bulletins = {}
        for mail_i in dct_1.keys():
            # 创建一个dict存储当前mail下af_zone与授权单号的关系
            _dct_btn = {}  # {'联通': ['08-01'], '电信': ['08-27']}
            for afk_i in dct_1[mail_i].keys():
                if not dct_1[mail_i][afk_i]['af_zone'] in _dct_btn.keys():
                    _dct_btn[dct_1[mail_i][afk_i]['af_zone']] = []
                _dct_btn[dct_1[mail_i][afk_i]['af_zone']].append(dct_1[mail_i][afk_i]['sbj_aid'])

            # 新建dict
            ki = list(dct_1[mail_i].keys())[0]   # 随机一个key即可
            dct_bulletins[mail_i] = {}
            dct_bulletins[mail_i]['eml_idf'] = dct_1[mail_i][ki]['eml_idf']
            dct_bulletins[mail_i]['ki'] = dct_1[mail_i][ki]['eml_idf']
            dct_bulletins[mail_i]['sbj'] = dct_1[mail_i][ki]['sbj']
            dct_bulletins[mail_i]['dir_attach'] = dct_1[mail_i][ki]['dir_attach']
            dct_bulletins[mail_i]['af_dict'] = _dct_btn

        dct_2 = {}  # 存放“{电信: [01-01, 01-02], }”
        for mail_i in dct_bulletins.keys():
            _dir_txt = dct_bulletins[mail_i]['dir_attach']
            _sbj = dct_bulletins[mail_i]['sbj']

            for az_i in dct_bulletins[mail_i]['af_dict'].keys():
                dct_2[az_i] = dct_bulletins[mail_i]['af_dict'][az_i]

                # 生成txt内容（带排序）
                lst_id = [str(ic[0]) + '-' + str(ic[-1]) for ic in
                          sorted([(int(ib[0]), int(ib[-1])) for ib in [ia.split('-') for ia in dct_2[az_i]]])]
                dct_bulletins[mail_i]['txt'] = '【' + az_i + '机房授权】' + '、'.join(lst_id) + '，' + _sbj

                # 处理【授权主题】内容
                # print('[debug] dct_bulletins[mail_i][txt]', dct_bulletins[mail_i]['txt'],[afi for afi in os.listdir(_dir_txt) if afi.endswith('.docx')])

                _cont = dct_bulletins[mail_i]['txt'].split('，')[-1]
                for i in _lst_ban_wd_:
                    _cont = _cont.replace(i, '')
                _txt = dct_bulletins[mail_i]['txt'].split('，')[0] + '，' + _cont

                _bat_fn1_ = _bat_fn_.replace('★', az_i)
                txt_cont = _bat_cmd_.replace('☆', dct_bulletins[mail_i]['txt'].replace('&', '^&'))
                with open(os.path.join(_dir_txt, _bat_fn1_), 'a') as f:
                    f.write(txt_cont)
                    # todo: 系统解决转义

                idf = dct_bulletins[mail_i]['eml_idf']
                if 'bulet' not in self.dict_log['ea'][idf].keys():
                    self.dict_log['ea'][idf]['bulet'] = {}
                self.dict_log['ea'][idf]['bulet'][az_i] = txt_cont

        # for log

        lst_idfs = sorted(list(set(dict_syn[k]['eml_idf'] for k in dict_syn.keys())))
        for idf in lst_idfs:
            d_tmp1 = self.dict_log['ea'][idf]['bulet']
            self.dict_log['ea'][idf]['bulets'] = '; \r\n'.join([': '.join((k, d_tmp1[k])) for k in sorted(d_tmp1.keys())])

            # [temp code] rename
            a1, a2 = os.path.split(self.dict_log['ea'][idf][0]['pth_eml'])
            _, c2 = os.path.splitext(a2)

            # print('s', self.dict_log['ea'][idf]['bulets'] + c2)
            # pprint.pprint(self.dict_log['ea'][idf])
            # print(os.path.join(r'C:\Users\7546671\Desktop\kiwigo_table\2_这个文件夹放【下载的邮件(.eml)】\机房邮件归档', self.dict_log['ea'][idf]['bulets'] + c2))

            f_n = self.dict_log['ea'][idf]['bulets']

            for _i in ['echo', ' |clip', ':']:
                while True:
                    if _i in f_n:
                        f_n = f_n.replace(_i, '')
                    else:
                        break

            # 机房邮件归档 （后续统一新建并识别目录）
            _dir_used_emls_ = os.path.join(_dir_eml_, '机房邮件归档')
            if not os.path.exists(_dir_used_emls_):
                os.makedirs(_dir_used_emls_)
            shutil.copyfile(src=self.dict_log['ea'][idf][0]['pth_eml'],
                            dst=os.path.join(_dir_used_emls_, f_n + c2))

    def _del_org_docx(self, dict_syn):
        for ki in dict_syn.keys():
            os.remove(os.path.join(dict_syn[ki]['dir_attach'], dict_syn[ki]['af_name'] + '.docx'))
            print('[删除文件] ', os.path.join(dict_syn[ki]['dir_attach'], dict_syn[ki]['af_name'] + '.docx'))
            shutil.copyfile(src=os.path.join(_dir_docx_, dict_syn[ki]['name_nw'] + '.docx'),
                            dst=os.path.join(dict_syn[ki]['dir_attach'], '#' + dict_syn[ki]['name_nw'] + '.docx'))

    def emls_to_doxs(self, _dct_ini_=_dct_ini_, sender_mode='origin'):
        _pth_doxs_arc = _dct_ini_['_dir_afbase_']
        _pth_eml = _dct_ini_['_dir_eml_']
        pth_out = _dct_ini_['_dir_out_']

        _dir_afb_kyzx_t_ = os.path.join(_pth_doxs_arc, '1.1_电信')
        _dir_afb_kyzx_u_ = os.path.join(_pth_doxs_arc, '1.4_联通')
        _dir_afb_rjzx_ = os.path.join(_pth_doxs_arc, '1.2_软件中心')
        _dir_afb_hlge_ = os.path.join(_pth_doxs_arc, '1.3_和林格尔')
        self.__dir_init(_dir=_pth_doxs_arc)
        self.__dir_init(_dir=_dir_afb_kyzx_t_)
        self.__dir_init(_dir=_dir_afb_kyzx_u_)
        self.__dir_init(_dir=_dir_afb_rjzx_)
        self.__dir_init(_dir=_dir_afb_hlge_)
        dict_eml = {}

        # 1. 遍历目录下eml文件
        for eml_i in os.listdir(_pth_eml):  # 仅遍历当前文件，不穿透深层文件夹
            if eml_i.endswith('.eml'):

                # 2. 处理eml
                with open(os.path.join(_pth_eml, eml_i), 'rb') as fhdl:
                    raw_email = fhdl.read()
                ep = eml_parser.EmlParser(include_attachment_data=True, include_raw_body=True)
                parsed_eml = ep.decode_email_bytes(raw_email)

                dict_eml['sbj'] = parsed_eml['header']['subject']
                dict_eml['pth_eml'] = os.path.join(_pth_eml, eml_i)
                # pprint.pprint(parsed_eml['header'])

                # 3. 保存信息
                eml_hd = parsed_eml['header']
                # 3.1 可选识别发件人（因申请人需转发leader后方可入授权邮箱，故需识别转发前的申请人）
                sender_ = ''.join([i.split(' <', 1)[0] for i in set(eml_hd['header']['from'])])
                if sender_mode == 'origin':
                    pass
                elif sender_mode == 'forward':
                    _m = parsed_eml['body'][0]['content']
                    sender_ = self.identify_sender(_mail_str=_m, )
                else:
                    print('[err] sender_mode should be "origin" or "forward", now, sender_mode = origin')

                date_ = datetime.datetime.strftime(eml_hd['date'], '%Y%m%d%H%M%S')
                eml_idf = date_ + '_' + sender_
                print(eml_idf + '\t' + dict_eml['sbj'])    # 打印正在处理哪封邮件

                tm.dicter(d0=self.dict_log['ea'], k1=eml_idf, k='date', v=date_)
                tm.dicter(d0=self.dict_log['ea'], k1=eml_idf, k='sender', v=sender_)
                print(self.dict_log['ea'])

                # 创建目录
                fname0 = date_ + '_' + sender_ + os.path.splitext(eml_i)[0]
                f_name = os.path.join(pth_out, fname0)
                if os.path.exists(f_name):
                    f_name = f_name + time.strftime('%H%M%S', time.localtime())
                if not os.path.exists(f_name):
                    os.makedirs(f_name)
                dict_eml['dir_attach'] = f_name

                # 4. 创建txt记录正文
                # todo: 后续根据邮件内容确认是否可以只取parsed_eml['body'][0]
                for i in parsed_eml['body']:
                    with open(os.path.join(f_name, _mail_txt_), 'a') as f:
                        f.write(i['content'])

                # 3. 保存eml附件
                if parsed_eml.get('attachment'):
                    for i in parsed_eml['attachment']:
                        x = base64.b64decode(i['raw'])
                        with open(os.path.join(f_name, i['filename']), 'wb') as f:
                            f.write(x)

                # 4. 用于同步信息
                dict_eml['sender'] = sender_
                dict_eml['date'] = date_
                dict_eml['eml_idf'] = eml_idf
                # self.dict_eml= copy.deepcopy(dict_eml)

                # 5. 处理授权单（.docx）
                dict_docx = self._idf_sqd(eml_idf=eml_idf, dir_eml_atm=f_name, dir_eml=os.path.join(_pth_eml, eml_i))
                if not -1 in dict_docx.keys():
                    self._syn_dict(dict_docx=dict_docx, dict_eml=dict_eml)

        # 6. 同步信息
        dict_syn = self._name_sqd(dict_syn=self.dict_syn)

        # todo: 7. 检查dict_syn是否均被识别
        lst_kyzx_t_k = [ki for ki in dict_syn.keys() if dict_syn[ki]['room'] in _lst_kyzx_t_]
        lst_kyzx_u_k = [ki for ki in dict_syn.keys() if dict_syn[ki]['room'] in _lst_kyzx_u_]
        lst_rjzx_k = [ki for ki in dict_syn.keys() if dict_syn[ki]['room'] in _lst_rjzx_]
        lst_hlge_k = [ki for ki in dict_syn.keys() if dict_syn[ki]['room'] in _lst_hlge_]
        self._dist_afid(dir_afb=_dir_afb_kyzx_t_, dict_syn=dict_syn, lst_k=lst_kyzx_t_k)
        self._dist_afid(dir_afb=_dir_afb_kyzx_u_, dict_syn=dict_syn, lst_k=lst_kyzx_u_k)
        self._dist_afid(dir_afb=_dir_afb_rjzx_, dict_syn=dict_syn, lst_k=lst_rjzx_k)
        self._dist_afid(dir_afb=_dir_afb_hlge_, dict_syn=dict_syn, lst_k=lst_hlge_k)

        self._txt(dict_syn=dict_syn)

        # 8. 删除原.docx授权文件（为简化邮件附件内容）
        self._del_org_docx(dict_syn=dict_syn)

        # 9. 删除.eml下载邮件
        # for eml_i in os.listdir(_pth_eml):  # 仅遍历当前文件，不穿透深层文件夹
        #     if eml_i.endswith('.eml'):
        #         pth_good_eml = os.path.join(_pth_eml, eml_i)
        #         if os.path.join(_pth_eml, eml_i) not in self.dict_log['er_eml']:
        #             os.remove(pth_good_eml)
        #             print('[删除文件] ', pth_good_eml)

    def doxs_to_xlx(self, ):
        self._doxs_to_xlx(_pth_dox=_dir_afb_kyzx_t_, _pth_xlx=_path_)
        self._doxs_to_xlx(_pth_dox=_dir_afb_kyzx_u_, _pth_xlx=_path_)
        self._doxs_to_xlx(_pth_dox=_dir_afb_rjzx_, _pth_xlx=_path_)
        self._doxs_to_xlx(_pth_dox=_dir_afb_hlge_, _pth_xlx=_path_)

    def _doxs_to_xlx(self, _pth_dox, _pth_xlx=_path_):
        # 1. 获取所有.docx文件
        authfile_list = []
        for _dir in os.listdir(_pth_dox):
            if _dir.endswith('.docx'):
                authfile_list.append(_dir)

        # 2. output
        list_bad = []
        lst_chgs = []  # 用于变更去重
        df0 = pd.DataFrame([], columns=_tab_hdlst_)
        for row_i, af_i in enumerate(authfile_list):
            afn_i, _ = os.path.splitext(af_i)

            # 3. 获取【序号】、【申请来源】
            df0.loc[row_i, '序号'] = row_i + 1
            df0.loc[row_i, '申请来源'] = afn_i

            # 4. 尝试获取【授权编号】、【机房】
            try:
                auth_id = re.match('^[\s]*[0-9]{6}-[0-9]+', afn_i).group()
                room = re.findall('[a-z]{1}[0-9]{4}', afn_i, re.I)[0]
                app_date = re.findall('[0-9]{4}-[0-9]{2}-[0-9]{2}', afn_i)[0]
                apper = re.findall('[\u4e00-\u9fa5]{2,3}', afn_i)[0]
                df0.loc[row_i, '机房授权编号'] = auth_id
                df0.loc[row_i, '机房'] = room
            except:
                df0.loc[row_i, '机房授权编号'] = ''
                df0.loc[row_i, '机房'] = ''
                list_bad.append(af_i)

            # 5. 对表格内容进行继续获取
            try:
                d0 = Document(os.path.join(_pth_dox, af_i))
                t_ = d0.tables  # 获取文件中的表格集
                t0 = t_[0]
            except:
                list_bad.append(af_i)
                continue

            # 6. 获取【人员所属】、【授权人数】、【设备】
            df0.loc[row_i, '人员所属'] = self.tm.fun_get_run_text(t1=t0.cell(2, 2)) + ';\n' + self.tm.fun_get_run_text(
                t1=t0.cell(1, 2))
            df0.loc[row_i, '授权人数'] = sum([int(i) for i in re.findall('[0-9]', self.tm.fun_get_run_text(t1=t0.cell(2, 2)))])

            str1 = self.tm.fun_get_run_text(t1=t0.cell(3, 2)) + '×' + self.tm.fun_get_run_text(t1=t0.cell(4, 2)) + ';' \
                   + self.tm.fun_get_run_text(t1=t0.cell(3, 4)) + '×' + self.tm.fun_get_run_text(t1=t0.cell(4, 4)) + ';' \
                   + self.tm.fun_get_run_text(t1=t0.cell(3, 7)) + '×' + self.tm.fun_get_run_text(t1=t0.cell(4, 7))

            str1 = str1.replace(';×', '')
            if str1.endswith('×'): str1 = str1.replace('×', '')
            if str1.endswith(';'): str1 = str1.replace(';', '')
            df0.loc[row_i, '设备'] = str1

            # 7. 获取【变更号】、【申请具体事项】
            chg_bare = self.tm.get_row_text(t1=t0, _row=6, _set='str', _mode='hard')
            chg_nosig = self.tm.rm_sig(_in=chg_bare, _rm=_lst_sig_)

            lst0 = []
            [lst0.extend(i) for i in [re.findall('CHG-\w{2,}-\w{2,}-\w{2,}-\d{8}-\d{4}', i) for i in chg_nosig]]

            df0.loc[row_i, '变更号'] = ', '.join(set(lst0))

            # 用于后续总变更去重
            lst_chg_i = list(set(lst0))
            if lst_chg_i:
                lst_chgs.extend(lst_chg_i)

            # 申请具体事项
            rsn_bare = self.tm.get_row_text(t1=t0, _row=7, _set='str', _mode='hard')

            lst_rsn_i = []
            for _r in rsn_bare:
                # 去除头尾坏字符
                for _sig in _lst_sig_:
                    _r = _r.strip(_sig)
                lst_rsn_i.append(_r)
            lst_rsn_i = self.tm.rm_sig(_in=lst_rsn_i, _rm=_lst_sig_ + _lst_docx_sig_)
            lst_rsn_i = list(set(lst_rsn_i))

            if len(lst_rsn_i) == 1:
                df0.loc[row_i, '申请具体事项'] = lst_rsn_i[0]
            elif len(lst_rsn_i) >= 2:
                df0.loc[row_i, '申请具体事项'] = lst_rsn_i[1]
            else:
                df0.loc[row_i, '申请具体事项'] = ''
                list_bad.append(af_i)

            # 8. 获取【进出时间】
            date_bare = self.tm.get_row_text(t1=t0, _row=8, _set='str', _mode='hard')
            date_bare = self.tm.rm_sig(_in=date_bare, _rm=_lst_sig_ + _lst_docx_sig_)
            lst_date_i = list(set(date_bare))

            # todo: 选最长的
            if len(lst_date_i) == 1:
                df0.loc[row_i, '进出时间'] = lst_date_i[0] + ' ~ ' + lst_date_i[0]
            elif len(lst_date_i) == 2:
                lst_date_i.sort()
                df0.loc[row_i, '进出时间'] = lst_date_i[0] + ' ~ ' + lst_date_i[1]
            else:
                df0.loc[row_i, '进出时间'] = ''
                list_bad.append(af_i)

        # 统计变更，变更去重
        set_chgs = set(lst_chgs)
        print('[总变更数（不去重）] ', len(lst_chgs))
        print('[总变更数（去重）] ', len(set_chgs))

        dct_chgs = {}
        for _chg in set_chgs:
            dct_chgs[_chg] = lst_chgs.count(_chg)
        print('[各变更次数（出现在授权单中几次）]', dct_chgs)
        df0.to_excel(os.path.join(_path_, os.path.basename(_pth_dox) + _bsname_xlsx_), index=False)

        list_bad = list(set(list_bad))
        print('[识别失败]\t', len(list_bad), ': \t', list_bad)

    # 用于从指定目录解压【压缩文件】到目标文件下
    def get_eml_once(self, dir_from, dir_to=_dct_ini_['_dir_eml_']):
        for _f in os.listdir(dir_from):
            p_f = os.path.join(dir_from, _f)

            if os.path.splitext(p_f)[-1] == '.eml':
                shutil.copyfile(src=p_f, dst=pth_f)
                os.remove(p_f)  # 删除源文件
                print('[删除文件] ', p_f)

            if os.path.splitext(p_f)[-1] in _lst_cps_:
                pth_f = os.path.join(dir_to, os.path.split(p_f)[-1])
                shutil.copyfile(src=p_f, dst=pth_f)

                # uncompress
                dir_tmp = os.path.join(dir_to, str(time.time()) + '_tmp')
                fb_code = tm.uncps(pth_f=pth_f, dir_ot=dir_tmp, )
                if fb_code == -1:   # 解压失败
                    continue
                if fb_code == 0:   # 解压成功
                    # 判断.eml
                    for _e in os.listdir(dir_tmp):
                        print(_e)
                        if _e.endswith('.eml'):
                            shutil.copyfile(src=os.path.join(dir_tmp, _e), dst=os.path.join(dir_to, _e))
                            print('[提取邮件] ', _e, '\tfrom\t', _f)
                    # os.remove(os.path.join(dir_to, _f))  # 删除【下载邮件目录】下的压缩包文件
                    # print('[删除文件] ', os.path.join(dir_to, _f))
                    # shutil.rmtree(dir_tmp)
                    # os.remove(p_f)   # 删除源文件
                    # print('[删除文件] ', p_f)

    # 回收log
    def show_log(self, ):
        # pprint.pprint(self.dict_log)
        df0 = pd.DataFrame([], columns=_lst_hd_log_ea)

        # 1. 处理邮件数
        self.dict_log['ea'].keys()
        # lst_emls = [self.dict_log['ea'][e] for e in self.dict_log['ea'].keys()]
        # print(lst_emls)
        row_i = 0
        for e in sorted(self.dict_log['ea'].keys()):
            ids = [k for k in self.dict_log['ea'][e].keys() if type(k) == int]   # 选出int为key的授权单dict
            for id in sorted(ids):
                # _lst_hd_log_ea = ['is_done', 'date', 'sender', 'dis_a_zone', 'dis_a_id', 'room', 'sbj', 'af/afs',
                #                   'af_org', 'af_edt', 'pth_eml', 'dir_eml_i']
                df0.loc[row_i, 'date'] = self.dict_log['ea'][e][id]['date']
                df0.loc[row_i, 'sender'] = self.dict_log['ea'][e][id]['sender']
                df0.loc[row_i, 'dis_a_zone'] = self.dict_log['ea'][e][id]['af_zone']
                df0.loc[row_i, 'dis_a_id'] = self.dict_log['ea'][e][id]['sbj_aid']
                df0.loc[row_i, 'room'] = self.dict_log['ea'][e][id]['room']
                df0.loc[row_i, 'sbj'] = self.dict_log['ea'][e][id]['sbj']
                df0.loc[row_i, 'af/afs'] = self.dict_log['ea'][e][id]['e_id']
                df0.loc[row_i, 'af_org'] = self.dict_log['ea'][e][id]['af_name']
                df0.loc[row_i, 'af_edt'] = self.dict_log['ea'][e][id]['name_nw']
                df0.loc[row_i, 'pth_eml'] = self.dict_log['ea'][e][id]['pth_eml']
                df0.loc[row_i, 'dir_eml_i'] = self.dict_log['ea'][e][id]['dir_attach']
                df0.loc[row_i, 'bulets'] = self.dict_log['ea'][e]['bulets']

                df0.loc[row_i, 'is_done'] = '√'
                row_i += 1
        df0.to_excel(os.path.join(_out_path_, 'res.xlsx'), index=None)

    '''可识别转发邮件发件人 （最新转发关系中，如收件人为王鹏，则发件人定为申请人） 输入为邮件正文str'''
    @classmethod
    def identify_sender(cls, _mail_str, _last_receiver='王鹏', unknow_name='~'):
        if (type(_last_receiver).__name__ == 'str'):
            lst_last_receiver = [_last_receiver]

        # 1. 提取[转发标识有效段落] flag
        lst_flag = re.compile('----邮件原文----.+?主题：', flags=re.S).findall(_mail_str)

        # 2 若receiver为王鹏 & sender不为王鹏；则sender为初始sender
        '''方法点'''
        # ↓ 此办法无法识别有多位收件人，或王鹏非第一位收件人的情况
        # _receiver = re.compile('收件人："(.+)"', re.M).findall(flg_sender[0])[0]
        _sender = unknow_name   # 若出错（识别失败），默认名
        try:
            for _flag in lst_flag:
                for _last_rcver in lst_last_receiver:
                    if _last_rcver in re.compile('收件人："(.+)"', re.M).findall(_flag):  # 有人发给[例：王鹏]的情况
                        _sender = re.compile('发件人："(.+)"', re.M).findall(_flag)[0]
                        return _sender   # case1: 识别成功，直接退出
            print('[err] case2 发件人未识别到，已置为', unknow_name)   # case2: lst_flag为空[]时运行
        except:
            print('[err] case3 发件人未识别到，已置为', unknow_name)   # case3: try中有任何报错时

        return _sender
